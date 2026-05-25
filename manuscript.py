"""
Manuscript generation: VTT transcript → Claude → polished Word document.
Intended for sermons, Bible teachings, and spoken-word content.
"""
import re
from pathlib import Path

import anthropic
from docx import Document
from docx.shared import Pt

SYSTEM_PROMPT = (
    "You are a manuscript editor specialising in converting spoken-word sermons and "
    "Bible teachings into polished, book-ready prose.\n\n"
    "Given a transcript of a spoken message, produce a clean manuscript:\n"
    "- Remove all filler words (um, uh, you know, like), false starts, and repetitions\n"
    "- Strip any timestamp markers or speaker labels\n"
    "- Restructure run-on oral sentences into clear written prose\n"
    "- Preserve all theological content, scripture references, and the speaker's voice faithfully\n"
    "- Organise into logical paragraphs; use ## headings for natural section breaks\n"
    "- Do not add content that wasn't in the original; do not summarise — this is a full manuscript\n\n"
    "Return ONLY the manuscript text. No preamble, no commentary, no markdown code fences."
)


def _vtt_to_text(vtt_content: str) -> str:
    """Extract plain text from a WebVTT string, stripping cue numbers and timestamps."""
    lines = []
    for line in vtt_content.splitlines():
        line = line.strip()
        if not line or line == "WEBVTT":
            continue
        if line.isdigit():
            continue
        if " --> " in line:
            continue
        lines.append(line)
    return " ".join(lines)


def _sanitize(text: str) -> str:
    return re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)


def _write_docx(title: str, manuscript_text: str, out_path: Path) -> None:
    """Write manuscript text to a Word document with Calibri styling."""
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(12)

    # Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(_sanitize(title))
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(18)
    title_run.bold = True

    doc.add_paragraph()

    for line in manuscript_text.splitlines():
        if line.startswith("## "):
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(_sanitize(line[3:].strip()))
            run.font.name = "Calibri"
            run.font.size = Pt(12)
            run.bold = True
        elif line.strip():
            p = doc.add_paragraph()
            run = p.add_run(_sanitize(line.strip()))
            run.font.name = "Calibri"
            run.font.size = Pt(12)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def vtt_to_manuscript(
    vtt_path: Path,
    title: str,
    docx_dest: Path,
    api_key: str,
    model: str = "claude-sonnet-4-6",
    status_callback=None,
) -> Path:
    """
    Convert a VTT transcript to a polished manuscript Word document.
    vtt_path: source .vtt file
    title: display title for the Word doc heading
    docx_dest: where to save the .docx
    Returns docx_dest.
    """
    if docx_dest.exists():
        if status_callback:
            status_callback("Using cached manuscript")
        return docx_dest

    vtt_content = vtt_path.read_text(encoding="utf-8")
    transcript_text = _vtt_to_text(vtt_content)

    if not transcript_text.strip():
        raise ValueError("VTT file produced no readable text")

    if status_callback:
        status_callback("Generating manuscript with Claude...")

    client = anthropic.Anthropic(api_key=api_key)
    try:
        message = client.messages.create(
            model=model,
            max_tokens=8192,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": f"TRANSCRIPT:\n{transcript_text}"}],
        )
        manuscript_text = message.content[0].text
    except anthropic.BadRequestError as exc:
        # Claude's content filter blocked the request — save the raw transcript instead
        if status_callback:
            status_callback("Claude blocked this content — saving raw transcript as manuscript...")
        manuscript_text = transcript_text

    _write_docx(title, manuscript_text, docx_dest)

    if status_callback:
        status_callback(f"Manuscript saved: {docx_dest.name}")

    return docx_dest
